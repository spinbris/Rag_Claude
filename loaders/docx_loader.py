
"""Microsoft Word document loader."""

from docx import Document
from typing import List, Dict
from .base import BaseLoader


class DocxLoader(BaseLoader):
    """Load content from Word (.docx) files."""
    
    def load(self, filepath: str) -> List[Dict[str, str]]:
        """
        Load and extract text from a Word document.
        
        Args:
            filepath: Path to .docx file
            
        Returns:
            List containing a single document dictionary
        """
        try:
            doc = Document(filepath)
            
            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = '\n\n'.join(paragraphs)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text += '\n' + row_text
            
            if not text.strip():
                raise ValueError("No text could be extracted from Word document")
            
            return [{
                'content': text,
                'source': filepath,
                'type': 'docx'
            }]
            
        except Exception as e:
            raise ValueError(f"Error loading Word document {filepath}: {str(e)}")

